
import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
from scipy import stats
import matplotlib.pyplot as plt
from datetime import datetime

st.set_page_config(page_title="App đo lường CLO", layout="wide")

st.title("📘 Ứng dụng đo lường Chuẩn đầu ra học phần (CLO)")
st.write("Tải lên file điểm (CSV/Excel). File có thể là điểm từng câu hỏi (mỗi cột Q1,Q2...) hoặc điểm tổng và cột phân bố câu hỏi.")

# ----------------- Upload dữ liệu -----------------
uploaded = st.file_uploader("1) Chọn file CSV/Excel (hỗ trợ .csv, .xls, .xlsx)", type=["csv","xls","xlsx"]) 

if uploaded is None:
    st.info("Vui lòng tải lên file dữ liệu để bắt đầu. Mẫu: MãSV, HoTen, Q1, Q2, ..., Qn hoặc MãSV, HoTen, DiemTong và file mapping Q->CLO.")
    st.stop()

# read file
try:
    if uploaded.name.endswith('.csv'):
        df = pd.read_csv(uploaded)
    else:
        df = pd.read_excel(uploaded)
except Exception as e:
    st.error(f"Không thể đọc file: {e}")
    st.stop()

st.subheader("Xem trước dữ liệu")
st.dataframe(df.head(5))

# auto-detect numeric columns
numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
all_cols = df.columns.tolist()

# Kiểm tra cột 'Tên học phần'
if 'Tên học phần' not in all_cols:
    st.error("❌ Dữ liệu chưa có cột 'Tên học phần'. Vui lòng kiểm tra lại file Excel/CSV.")
    st.stop()

# --- Chọn học phần cần làm việc ---
st.subheader("Chọn học phần cần đo lường")
hocphan_list = df['Tên học phần'].dropna().unique().tolist()
selected_hocphan = st.selectbox("🎓 Chọn học phần bạn cần đo", hocphan_list)

# --- Lọc dữ liệu theo học phần được chọn ---
df_hp = df[df['Tên học phần'] == selected_hocphan].copy()

# --- Loại các cột không phải điểm ---
ignore_cols = ['Tên học phần', 'IDSV', 'Họ và tên SV', 'Lớp', 'Số phách', 'Tổng điểm', 'Mã đề']

# --- Lấy danh sách cột điểm thực sự có dữ liệu ---
numeric_cols = []
for c in df_hp.columns:
    if c not in ignore_cols:
        try:
            # Kiểm tra nếu cột có ít nhất một giá trị số (không NaN)
            if pd.to_numeric(df_hp[c], errors='coerce').notna().sum() > 0:
                numeric_cols.append(c)
        except:
            pass

if len(numeric_cols) == 0:
    st.warning("⚠️ Không tìm thấy cột điểm hợp lệ nào trong dữ liệu học phần này.")
    st.stop()

# --- Khai báo điểm tối đa cho từng câu hỏi ---
st.header(f"🧮 Khai báo điểm tối đa cho từng câu hỏi - {selected_hocphan}")
st.write("Nhập điểm tối đa cho từng câu hỏi (chỉ hiển thị các câu có dữ liệu thực tế):")

max_scores = {}

for q in numeric_cols:
    c1, c2 = st.columns([1, 1])
    with c1:
        st.write(f"**{q}**")
    with c2:
        max_scores[q] = st.number_input(
            label="",
            min_value=0.1,
            max_value=10.0,
            value=1.0,
            step=0.1,
            key=f"max_{selected_hocphan}_{q}"
        )

st.success(f"✅ Đã khai báo điểm tối đa cho {len(numeric_cols)} câu hỏi có dữ liệu của học phần {selected_hocphan}.")

# ------------------ KHAI BÁO CHUẨN ĐẦU RA (CĐR) ------------------
st.header(f"🎯 Khai báo Chuẩn đầu ra (CĐR) cho học phần {selected_hocphan}")

st.subheader("Nhập số lượng, nội dung và câu hỏi cần mapping")

num_cdr = st.number_input("Số lượng CĐR của học phần", min_value=1, max_value=20, step=1, value=1)

cdr_data = []  # danh sách lưu kết quả khai báo

for i in range(int(num_cdr)):
    st.markdown(f"### 🧩 Chuẩn đầu ra {i+1}")

    with st.container():
        col1, col2, col3, col4, col5, col6 = st.columns([1, 3, 1, 1, 1, 1])

        # --- 1. Tên viết tắt ---
        with col1:
            ten_cdr = st.text_input(
                "Tên viết tắt",
                value=f"CĐR{i+1}",
                key=f"cdr_name_{selected_hocphan}_{i}"
            )

        # --- 2. Nội dung CĐR ---
        with col2:
            noidung_cdr = st.text_area(
                "Nội dung CĐR",
                placeholder="Nhập mô tả nội dung CĐR...",
                key=f"cdr_content_{selected_hocphan}_{i}",
                height=80
            )

        # --- 3. Câu hỏi thuộc CĐR ---
        with col3:
            cauhoi_chon = st.multiselect(
                "Dữ liệu được lấy từ",
                options=numeric_cols,
                placeholder="Chọn câu hỏi...",
                key=f"cdr_questions_{selected_hocphan}_{i}",
                help="Chọn các câu hỏi được dùng để đo lường CĐR này"
            )

        # --- 4. Tỷ lệ điểm tối thiểu (%) ---
        with col4:
            tile_diemtoithieu = st.number_input(
                "Tỷ lệ điểm tối thiểu (%)",
                min_value=0.0, max_value=100.0, step=5.0, value=40.0,
                key=f"cdr_tilemin_{selected_hocphan}_{i}"
            )

        # --- 5. Điểm tối thiểu (tự động tính) ---
        with col5:
            if cauhoi_chon:
                diem_tb_max = np.mean([max_scores[q] for q in cauhoi_chon])
                diem_toithieu = diem_tb_max * (tile_diemtoithieu / 100)
            else:
                diem_toithieu = 0.0
            st.number_input(
                "Điểm tối thiểu",
                value=round(diem_toithieu, 2),
                disabled=True,
                key=f"cdr_diemmin_{selected_hocphan}_{i}"
            )

        # --- 6. Tỷ lệ SV đạt kỳ vọng (%) ---
        with col6:
            tile_kyvong = st.number_input(
                "Tỷ lệ kỳ vọng (%)",
                min_value=0.0, max_value=100.0, step=5.0, value=75.0,
                key=f"cdr_tilekyvong_{selected_hocphan}_{i}"
            )

    # Thêm đường kẻ phân cách
    st.markdown("---")


    # --- Lưu vào danh sách ---
    cdr_data.append({
        "Tên CĐR": ten_cdr,
        "Nội dung": noidung_cdr,
        "Câu hỏi": ", ".join(cauhoi_chon) if cauhoi_chon else "",
        "Tỷ lệ điểm tối thiểu (%)": tile_diemtoithieu,
        "Điểm tối thiểu": round(diem_toithieu, 2),
        "Tỷ lệ kỳ vọng (%)": tile_kyvong
    })

# ------------------ HIỂN THỊ BẢNG KHAI BÁO CĐR ------------------
if cdr_data:
    st.subheader("📋 Tổng hợp thông tin CĐR đã khai báo")
    df_cdr = pd.DataFrame(cdr_data)
    st.dataframe(df_cdr, use_container_width=True)
    # 👉 Lưu vào session để dùng khi xuất báo cáo
    st.session_state.df_cdr = df_cdr

    # --- Nút hoàn tất khai báo ---
    if st.button("✅ Hoàn tất khai báo CĐR"):
        # Lưu dữ liệu vào session_state để dùng cho các phần sau
        st.session_state['df_cdr'] = df_cdr
        st.success(f"✅ Đã hoàn tất khai báo {len(df_cdr)} Chuẩn đầu ra (CĐR) cho học phần **{selected_hocphan}**.")
        st.balloons()  # hiệu ứng vui mắt khi hoàn tất

    # ✅ Tạo biến cdr_mapping để dùng cho các phần sau
    cdr_mapping = {}
    for i, row in df_cdr.iterrows():
        key = row["Tên CĐR"]
        cdr_mapping[key] = {
            "Nội dung": row["Nội dung"],
            "Câu hỏi": [q.strip() for q in row["Câu hỏi"].split(",") if q.strip()],
            "Điểm tối thiểu": row["Điểm tối thiểu"],
            "Tỷ lệ điểm tối thiểu (%)": row["Tỷ lệ điểm tối thiểu (%)"],
            "Tỷ lệ kỳ vọng (%)": row["Tỷ lệ kỳ vọng (%)"]
        }

    # Lưu vào session để phần sau dùng được
    st.session_state["cdr_mapping"] = cdr_mapping

# ------------------ PHÂN TÍCH KẾT QUẢ ĐẠT CĐR ------------------
st.header("📊 Phân tích thống kê kết quả đạt Chuẩn đầu ra (CĐR)")

if 'df_cdr' not in locals() or df_cdr.empty:
    st.warning("⚠️ Chưa có dữ liệu khai báo CĐR để phân tích.")
else:
    results = []

    tong_sv = len(df_hp)

    for idx, row in df_cdr.iterrows():
        cdr_name = row["Tên CĐR"]
        noi_dung = row["Nội dung"]
        cauhoi_list = [q.strip() for q in row["Câu hỏi"].split(",") if q.strip() in df_hp.columns]

        if not cauhoi_list:
            results.append({
                "CĐR": cdr_name,
                "Nội dung": noi_dung,
                "Điểm tối đa CĐR": "-",
                "Điểm tối thiểu đạt CĐR": "-",
                "Tổng SV đạt": "-",
                "Tỷ lệ SV đạt (%)": "-",
                "Kết quả": "-"
            })
            continue

        # 1. Điểm tối đa CĐR = tổng điểm tối đa các câu hỏi
        diem_toi_da_cdr = sum([max_scores.get(q, 0) for q in cauhoi_list])

        # 2. Điểm tối thiểu đạt CĐR (tính theo % nếu có nhiều câu hỏi)
        diem_toi_thieu_cdr = row["Điểm tối thiểu"]
        if len(cauhoi_list) > 1:
            diem_toi_thieu_cdr = sum([max_scores[q] * (row["Tỷ lệ điểm tối thiểu (%)"] / 100) for q in cauhoi_list])

        # 3. Tổng điểm thực tế sinh viên theo CĐR
        df_hp["Tổng_" + cdr_name] = df_hp[cauhoi_list].sum(axis=1, skipna=True)

        # 4. Tính số SV đạt và tỷ lệ đạt
        sv_dat = (df_hp["Tổng_" + cdr_name] >= diem_toi_thieu_cdr).sum()
        tyle_dat = round((sv_dat / tong_sv) * 100, 2) if tong_sv > 0 else 0
        
        # 5. Kết quả đo lường CĐR
        nhan_xet = "ĐẠT ✅" if tyle_dat >= tile_kyvong else "KHÔNG ĐẠT ❌"

        # Thêm vào danh sách kết quả
        results.append({
            "CĐR": cdr_name,
            "Nội dung": noi_dung,
            "Điểm tối đa CĐR": round(diem_toi_da_cdr, 2),
            "Điểm tối thiểu đạt CĐR": round(diem_toi_thieu_cdr, 2),
            "Tổng SV đạt": int(sv_dat),
            "Tỷ lệ SV đạt (%)": tyle_dat,
            "Kết quả": nhan_xet,
        })

    # Tạo DataFrame kết quả
    df_thongke = pd.DataFrame(results)
    df_thongke.index = np.arange(1, len(df_thongke) + 1)
    df_thongke.reset_index(inplace=True)
    df_thongke.rename(columns={"index": "TT"}, inplace=True)

    st.dataframe(df_thongke, use_container_width=True)

# 👉 Lưu kết quả thống kê đạt CĐR
    st.session_state.df_thongke = df_thongke


# ------------------ BIỂU ĐỒ TỶ LỆ SV ĐẠT CĐR ------------------
st.subheader("📊 Biểu đồ tỷ lệ sinh viên đạt Chuẩn đầu ra (CĐR) so với tỷ lệ kỳ vọng")

try:
    # Lấy dữ liệu từ bảng thống kê
    cdr_labels = df_thongke["CĐR"].tolist()
    ty_le_dat = df_thongke["Tỷ lệ SV đạt (%)"].tolist()

    # Lấy tỷ lệ kỳ vọng từ df_cdr (dựa trên cùng thứ tự CĐR)
    ty_le_ky_vong = []
    for cdr in cdr_labels:
        row = df_cdr[df_cdr["Tên CĐR"] == cdr]
        if not row.empty:
            ty_le_ky_vong.append(float(row["Tỷ lệ kỳ vọng (%)"].values[0]))
        else:
            ty_le_ky_vong.append(0)

    # Vẽ biểu đồ
    fig_tyle_cdr, ax = plt.subplots(figsize=(10, 5))

    # Cột tỷ lệ đạt
    bars = ax.bar(cdr_labels, ty_le_dat, color="#4CAF50", alpha=0.85, label="Tỷ lệ SV đạt (%)")

    # Dòng tỷ lệ kỳ vọng
    ax.plot(cdr_labels, ty_le_ky_vong, color="orange", marker="o", linewidth=2, label="Tỷ lệ kỳ vọng (%)")

    # Hiển thị giá trị phần trăm trên đầu cột
    for bar in bars:
        height = bar.get_height()
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            height + 1,
            f"{height:.1f}%",
            ha="center",
            va="bottom",
            fontsize=10,
            fontweight="bold"
        )

    # Thiết lập nhãn
    ax.set_xlabel("Chuẩn đầu ra (CĐR)", fontsize=11)
    ax.set_ylabel("Tỷ lệ sinh viên đạt (%)", fontsize=11)
    ax.set_ylim(0, 110)
    ax.set_title(f"Tỷ lệ sinh viên đạt CĐR – {selected_hocphan}", fontsize=13, fontweight="bold")

    # Di chuyển chú giải ra ngoài để không che dữ liệu
    ax.legend(loc="upper left", bbox_to_anchor=(1, 1))
    ax.grid(axis='y', linestyle='--', alpha=0.7)

    # Hiển thị trên giao diện
    st.pyplot(fig_tyle_cdr)

    # 👉 Lưu dữ liệu và biểu đồ vào session để xuất Word
    st.session_state.df_thongke = df_thongke
    st.session_state.fig_tyle_cdr = fig_tyle_cdr

    st.success("✅ Biểu đồ tỷ lệ SV đạt CĐR đã được tạo và lưu thành công!")

except Exception as e:
    st.error(f"⚠️ Lỗi khi tạo biểu đồ: {e}")


    # Xuất Excel
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        df_thongke.to_excel(writer, index=False, sheet_name='ThongKe_CDR')
        worksheet = writer.sheets['ThongKe_CDR']
        worksheet.set_column('A:A', 5)   # TT
        worksheet.set_column('B:B', 10)  # CĐR
        worksheet.set_column('C:C', 60)  # Nội dung CĐR
        worksheet.set_column('D:F', 18)  # Các cột điểm và SV
    st.download_button(
        label="📥 Tải bảng thống kê CĐR (Excel)",
        data=buffer.getvalue(),
        file_name=f"ThongKe_CDR_{selected_hocphan}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    
# =====================================================

# 📊 BẢNG PHÂN LOẠI NGƯỜI HỌC ĐẠT CĐR
# =====================================================
st.subheader(f"📊 Thống kê phân loại số lượng người học đạt CĐR – {selected_hocphan}")

# Tổng số SV có điểm
total_sv = len(df_hp)

# Tạo dataframe lưu kết quả
rows = []

import re

# 1) Tạo danh sách CĐR (dựa vào df_cdr nếu có, else dùng question_to_clo)
cdr_rows = []
if 'df_cdr' in locals() and not df_cdr.empty:
    for _, r in df_cdr.iterrows():
        cdr_rows.append(r.to_dict())
elif 'question_to_clo' in locals() and question_to_clo:
    inv = {}
    for q, clo in question_to_clo.items():
        inv.setdefault(clo, []).append(q)
    for clo, qlist in inv.items():
        cdr_rows.append({
            "Tên CĐR": clo,
            "Nội dung": "",
            "Câu hỏi": ", ".join(qlist),
            "Tỷ lệ điểm tối thiểu (%)": 0,
            "Điểm tối thiểu": 0
        })
else:
    st.error("Không tìm thấy thông tin CĐR. Vui lòng khai báo CĐR trước khi chạy phần phân loại.")
    st.stop()

# 2) Chuẩn bị bảng kết quả phân loại
rows = []

for r in cdr_rows:
    cdr_key = r.get('Tên CĐR') or r.get('CĐR') or r.get('CLO') 
    if not cdr_key:
        continue

    raw_q = r.get('Câu hỏi', '')
    if isinstance(raw_q, str):
        cau_hoi = [q.strip() for q in raw_q.split(',') if q.strip()]
    elif isinstance(raw_q, (list, tuple, np.ndarray)):
        cau_hoi = [q for q in raw_q]
    else:
        cau_hoi = []

    cau_hoi = [q for q in cau_hoi if q in df_hp.columns]

    if len(cau_hoi) == 0:
        rows.append({
            "Ký hiệu CĐR": cdr_key,
            "Tổng số SV": total_sv,
            "Loại A (Đạt)": 0, "Loại B (Đạt)": 0, "Loại C (Đạt)": 0, "Loại D (Đạt)": 0, "Loại F (Không đạt)": 0,
            "Tỷ lệ A (Đạt) (%)": 0, "Tỷ lệ B (Đạt) (%)": 0, "Tỷ lệ C (Đạt) (%)": 0, "Tỷ lệ D (Đạt) (%)": 0, "Tỷ lệ F (Không đạt) (%)": 0
        })
        continue

    diem_toi_da_cdr = sum([max_scores.get(q, 0) for q in cau_hoi])

    if 'Điểm tối thiểu' in r and r.get('Điểm tối thiểu') not in [None, '', np.nan]:
        diem_toi_thieu_cdr = float(r.get('Điểm tối thiểu'))
    else:
        ratio = float(r.get('Tỷ lệ điểm tối thiểu (%)') or 0) / 100.0
        diem_toi_thieu_cdr = round(diem_toi_da_cdr * ratio, 6)

    safe = re.sub(r'[^0-9a-zA-Z_]', '_', str(cdr_key))
    sum_col = f"__sum_{safe}"
    qd_col = f"__qd_{safe}"

    df_hp[sum_col] = df_hp[cau_hoi].apply(pd.to_numeric, errors='coerce').fillna(0).sum(axis=1)

    if diem_toi_da_cdr > 0:
        df_hp[qd_col] = (df_hp[sum_col] / diem_toi_da_cdr) * 10
    else:
        df_hp[qd_col] = 0.0

    A = ((df_hp[qd_col] >= 8.5) & (df_hp[qd_col] <= 10)).sum()
    B = ((df_hp[qd_col] >= 7.0) & (df_hp[qd_col] < 8.5)).sum()
    C = ((df_hp[qd_col] >= 5.5) & (df_hp[qd_col] < 7.0)).sum()
    D = ((df_hp[qd_col] >= 4.0) & (df_hp[qd_col] < 5.5)).sum()
    F = (df_hp[qd_col] < 4.0).sum()

    pct = lambda x: round((x / total_sv) * 100, 2) if total_sv > 0 else 0.0
    A_pct, B_pct, C_pct, D_pct, F_pct = map(pct, [A, B, C, D, F])

    rows.append({
        "Ký hiệu CĐR": cdr_key,
        "Tổng số SV": total_sv,
        "Loại A (Đạt)": int(A), "Loại B (Đạt)": int(B), "Loại C (Đạt)": int(C), "Loại D (Đạt)": int(D), "Loại F (Không đạt)": int(F),
        "Tỷ lệ A (Đạt) (%)": A_pct, "Tỷ lệ B (Đạt) (%)": B_pct, "Tỷ lệ C (Đạt) (%)": C_pct, "Tỷ lệ D (Đạt) (%)": D_pct, "Tỷ lệ F (Không đạt) (%)": F_pct
    })

# 3️⃣ Hiển thị bảng kết quả
df_phanloai = pd.DataFrame(rows)

if df_phanloai.empty:
    st.warning("Không có kết quả phân loại CĐR để hiển thị.")
else:
    def color_val(val, is_fail=False):
        if is_fail and val > 0:
            return 'background-color: #ff9999'
        elif not is_fail and val > 0:
            return 'background-color: #b3ffb3'
        return ''

    styled = df_phanloai.style.applymap(lambda v: color_val(v, False),
                                        subset=['Loại A (Đạt)','Loại B (Đạt)','Loại C (Đạt)','Loại D (Đạt)']) \
                              .applymap(lambda v: color_val(v, True),
                                        subset=['Loại F (Không đạt)']) \
                              .format({c: '{:.2f}' for c in ['Tỷ lệ A (Đạt) (%)','Tỷ lệ B (Đạt) (%)','Tỷ lệ C (Đạt) (%)','Tỷ lệ D (Đạt) (%)','Tỷ lệ F (Không đạt) (%)']})

    st.subheader("📋 Bảng phân loại A-B-C-D-F theo CĐR (số lượng & tỷ lệ)")
    st.write(styled)

    # ✅ Lưu DataFrame vào session_state để xuất Word
    st.session_state.df_af_summary = df_phanloai.copy()

    # --- Xuất Excel ---
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        df_phanloai.to_excel(writer, index=False, sheet_name='PhanLoai_CDR')
        ws = writer.sheets['PhanLoai_CDR']
        ws.set_column('A:A', 14)
        ws.set_column('B:B', 12)
        ws.set_column('C:G', 10)
        ws.set_column('H:L', 14)
    st.download_button("📥 Tải phân loại CĐR (Excel)", data=buffer.getvalue(),
                       file_name=f"PhanLoai_CDR_{selected_hocphan}.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # 4️⃣ BIỂU ĐỒ stacked bar (tỷ lệ % A..F)
st.subheader("🎨 Biểu đồ phân bố A–B–C–D–F theo CĐR")

import matplotlib.pyplot as plt

# Cấu hình loại và màu
categories = [
    "Tỷ lệ A (Đạt) (%)",
    "Tỷ lệ B (Đạt) (%)",
    "Tỷ lệ C (Đạt) (%)",
    "Tỷ lệ D (Đạt) (%)",
    "Tỷ lệ F (Không đạt) (%)"
]
colors = ['#2ca02c', '#98df8a', '#c7e9b4', '#ffe680', '#ff6666']

try:
    # Vẽ stacked bar chart
    fig_af_chart, ax = plt.subplots(figsize=(10, 6))
    bottom = np.zeros(len(df_phanloai))

    for i, cat in enumerate(categories):
        vals = df_phanloai[cat].values
        ax.bar(df_phanloai["Ký hiệu CĐR"], vals, bottom=bottom, color=colors[i], label=cat)
        bottom += vals

    # Nhãn trục và tiêu đề
    ax.set_ylabel("Tỷ lệ (%)", fontsize=11)
    ax.set_ylim(0, 100)
    ax.set_xlabel("Ký hiệu CĐR", fontsize=11)
    ax.set_title(f"Phân bố A–B–C–D–F theo CĐR – {selected_hocphan}", fontsize=13, fontweight="bold")

    # ✅ Di chuyển chú giải ra ngoài để không che cột
    ax.legend(
        title="Phân loại",
        loc="upper left",
        bbox_to_anchor=(1.02, 1),
        borderaxespad=0,
        fontsize=10,
        title_fontsize=11
    )

    # ✅ Hiển thị giá trị phần trăm trong cột
    for idx, row in df_phanloai.iterrows():
        cum = 0
        for cat in categories:
            val = row[cat]
            if val >= 3:  # chỉ hiển thị nếu đủ lớn
                ax.text(idx, cum + val / 2, f"{val:.1f}%", ha='center', va='center', fontsize=9)
            cum += val

    # Thêm lưới ngang để dễ đọc
    ax.grid(axis='y', linestyle='--', alpha=0.7)

    # Hiển thị biểu đồ
    st.pyplot(fig_af_chart)

    # ✅ Lưu biểu đồ và bảng phân loại vào session_state để xuất Word
    st.session_state.fig_af_chart = fig_af_chart
    st.session_state.df_phanloai = df_phanloai

    st.success("✅ Biểu đồ phân loại A–F đã được tạo và lưu thành công!")

except Exception as e:
    st.error(f"⚠️ Lỗi khi tạo biểu đồ A–F: {e}")


# ===================== 📄 XUẤT BÁO CÁO CLO (WORD) =====================
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import os

st.subheader("📘 Xuất báo cáo CLO (Word)")

# --- Nhập nhận xét & đề xuất ---
if "nhanxet" not in st.session_state:
    st.session_state.nhanxet = ""
if "dexuat" not in st.session_state:
    st.session_state.dexuat = ""

st.session_state.nhanxet = st.text_area("✍️ Nhập nhận xét tổng quan:", value=st.session_state.nhanxet, key="nhanxet_text")
st.session_state.dexuat = st.text_area("💡 Nhập đề xuất cải tiến:", value=st.session_state.dexuat, key="dexuat_text")

# --- Nút tạo báo cáo ---
if st.button("📤 Tạo báo cáo CLO (Word)", key="btn_export_word"):
    try:
        st.info("🧾 Đang tạo báo cáo... Vui lòng chờ trong giây lát.")
        doc = Document()

        # Cài đặt style font
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        style.font.size = Pt(12)

        # ==================== PHẦN I ====================
        doc.add_heading("PHẦN I. THÔNG TIN CHUNG", level=1)

        doc.add_paragraph(f"📘 Tên học phần: {selected_hocphan}")
        doc.add_paragraph(f"👨‍🎓 Số lượng sinh viên: {len(df_hp)}")

        if "df_cdr" in st.session_state:
            df_cdr = st.session_state.df_cdr
        else:
            df_cdr = df_cdr

        doc.add_paragraph("Tổng hợp thông tin CĐR đã khai báo:")
        table1 = doc.add_table(rows=1, cols=len(df_cdr.columns))
        table1.style = 'Table Grid'
        hdr_cells = table1.rows[0].cells
        for j, col_name in enumerate(df_cdr.columns):
            p = hdr_cells[j].paragraphs[0]
            run = p.add_run(col_name)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for _, row in df_cdr.iterrows():
            row_cells = table1.add_row().cells
            for j, value in enumerate(row):
                row_cells[j].text = str(value)

        # ==================== PHẦN II ====================
        doc.add_heading("PHẦN II. PHÂN TÍCH THỐNG KÊ KẾT QUẢ ĐẠT CHUẨN ĐẦU RA", level=1)

        df_thongke = st.session_state.get("df_thongke")
        fig_tyle_cdr = st.session_state.get("fig_tyle_cdr")

        # --- Bảng thống kê mức độ đạt ---
        doc.add_paragraph("1️⃣ Bảng thống kê mức độ đạt CĐR:")
        if df_thongke is not None:
            table2 = doc.add_table(rows=1, cols=len(df_thongke.columns))
            table2.style = 'Table Grid'
            hdr = table2.rows[0].cells
            for j, col in enumerate(df_thongke.columns):
                run = hdr[j].paragraphs[0].add_run(col)
                run.bold = True
                hdr[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for _, row in df_thongke.iterrows():
                cells = table2.add_row().cells
                for j, val in enumerate(row):
                    cells[j].text = str(val)
        else:
            doc.add_paragraph("⚠️ Chưa có dữ liệu thống kê mức độ đạt CĐR.")

        # --- Biểu đồ tỷ lệ sinh viên đạt CĐR ---
        doc.add_paragraph("2️⃣ Biểu đồ tỷ lệ sinh viên đạt Chuẩn đầu ra (CĐR) so với tỷ lệ kỳ vọng:")
        if fig_tyle_cdr is not None:
            chart_path = "chart_tylecdr.png"
            fig_tyle_cdr.savefig(chart_path, bbox_inches="tight")
            doc.add_picture(chart_path, width=Inches(6))
            os.remove(chart_path)
        else:
            doc.add_paragraph("⚠️ Không thể chèn biểu đồ tỷ lệ đạt CĐR.")

        # ==================== PHẦN III ====================
        doc.add_heading("PHẦN III. THỐNG KÊ PHÂN LOẠI NGƯỜI HỌC ĐẠT CĐR", level=1)

        df_af_summary = st.session_state.get("df_af_summary")
        fig_af_chart = st.session_state.get("fig_af_chart")

        # --- Bảng phân loại A-F ---
        doc.add_paragraph("1️⃣ Bảng phân loại A–B–C–D–F theo CĐR (số lượng & tỷ lệ):")
        if df_af_summary is not None:
            table3 = doc.add_table(rows=1, cols=len(df_af_summary.columns))
            table3.style = 'Table Grid'
            hdr3 = table3.rows[0].cells
            for j, col in enumerate(df_af_summary.columns):
                run = hdr3[j].paragraphs[0].add_run(col)
                run.bold = True
                hdr3[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for _, row in df_af_summary.iterrows():
                cells = table3.add_row().cells
                for j, val in enumerate(row):
                    cells[j].text = str(val)
        else:
            doc.add_paragraph("⚠️ Chưa có dữ liệu phân loại A–F để hiển thị.")

        # --- Biểu đồ phân bố điểm A-F ---
        doc.add_paragraph("2️⃣ Biểu đồ phân bố điểm A–B–C–D–F theo CĐR:")
        if fig_af_chart is not None:
            chart_path2 = "chart_af.png"
            fig_af_chart.savefig(chart_path2, bbox_inches="tight")
            doc.add_picture(chart_path2, width=Inches(6))
            os.remove(chart_path2)
        else:
            doc.add_paragraph("⚠️ Không thể chèn biểu đồ A–F.")

        # ==================== PHẦN IV ====================
        doc.add_heading("PHẦN IV. NHẬN XÉT – ĐỀ XUẤT", level=1)
        nhanxet = st.session_state.nhanxet
        dexuat = st.session_state.dexuat

        doc.add_paragraph(f"1️⃣ Nhận xét: {nhanxet}")
        doc.add_paragraph(f"2️⃣ Đề xuất: {dexuat}")

        # ==================== LƯU & TẢI ====================
        output_path = f"Bao_cao_CLO_{selected_hocphan}.docx"
        doc.save(output_path)

        with open(output_path, "rb") as f:
            st.download_button(
                label="📥 Tải xuống báo cáo Word (A4)",
                data=f,
                file_name=output_path,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.success("✅ Báo cáo Word đã được tạo thành công!")

    except Exception as e:
        st.error(f"⚠️ Lỗi khi tạo báo cáo: {e}")


